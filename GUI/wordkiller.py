# -*- coding: utf-8 -*-
"""
Created on Wed Sep 21 16:08:06 2022

@author: LV
"""

import docx
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn
import xlrd
import numpy as np
import pandas as pd

# df 传入要求是已经处理好的结果 path 为文件生成的位置
def make(df,start,end,path):

    document=docx.Document()
    namelist=list(df['企业名称'])
    sumlist=list(df['纳税合计'])
    sumlist2=list(df['Unnamed: 32'])


    def convert(num):
        ch_num=['零','一','二','三','四','五','六','七','八','九','十']
        s_unit=['','十','百','千']
        b_unit=['','万', '亿', '兆', '京', '垓', '秭', '穣', '沟', '涧', '正', '载', '极', ['恒河沙'], ['阿僧祇'], ['那由他'], ['不可思议'], ['无量大数']]
        numlist = list(map(int,str(num)))
        numlist.reverse()
        l=[]
        j=0
        for i in range(0, len(numlist), 4):
            p=[]
            if sum(numlist[i:i+4]):
                for ii in range(0, 4):
                    if (i+ii)<len(numlist):
                        if numlist[i+ii]:
                            p.append([s_unit[ii],ch_num[numlist[i+ii]]])
                        else:
                            if p and p[-1] != ['零']:
                                p.append(['零'])
                if j < len(b_unit):
                    l.append([b_unit[j]]+p)
                else:
                    return "数字太大,超出计量范围！！"
            else:
                if l and l[-1]!=['零'] and l[-1][-1]!=['零']:
                    l.append(['零'])
            j += 1
        l = [c for a in l for b in a for c in b if c]
        l.reverse()
        l=''.join(l)
        return l    
    a=1
    for i in range(len(namelist)):    

        
        

        a1=convert(a)
        a+=1
        b=namelist[i]    
        c= '      '+str(start)[:4]+'年'+str(start)[5:6]+'月至'+str(end)[:4]+'年'+str(end)[4:7]+'月总纳税(入库期)'
        d1=sumlist[i]
        d="{:,.2f}".format(d1)
        e='元，可申请产业扶持财政奖励合计'
        f1=sumlist2[i]
        f="{:,.2f}".format(f1)
        document.add_heading(a1+'、'+b,level = 1)
        document.styles['Normal'].font.size = Pt(12)
        document.styles['Normal'].font.name = u'仿宋_GB2312'
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')   
        p = document.add_paragraph(style='Normal')
        p2 = document.add_paragraph(style='Normal')
        run = p2.add_run(c+d+e+f+'元')
    document.save(path)
    print('《运营总部关于申请财政奖励的说明》已生成')

