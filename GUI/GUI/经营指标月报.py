# -*- coding: utf-8 -*-
"""
Created on Mon Apr 11 10:16:04 2022

@author: EDY
"""

import docx
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn
import xlrd
import numpy as np
import pandas as pd
import datetime 
if __name__=='__main__':

    def makereport(start=202201,end=202206):    
        xlsx = pd.ExcelFile('d:/DATABASE/budget/业务预算-招商服务09.xlsx')
        df = pd.read_excel(xlsx, '业务实际-汇总')
        df1 = pd.read_excel(xlsx, '业务实际-商务')
        df2 = pd.read_excel(xlsx, '业务实际-化工')
        df3 = pd.read_excel(xlsx, '业务实际-市场开发')
        df4 = pd.read_excel(xlsx, '收入明细表',header=1)
        df5 = pd.read_excel(xlsx, '业务预算-商务')
        df6 = pd.read_excel('d:/DATABASE/all/客户开票.xlsx')
        document=docx.Document()
        year='2022'
        
        month=int(str(start)[4:6]) #输入需要报告的开始月份
        endmonth=int(str(end)[4:6]) #输入需要报告的截至月份
        start_date=start
        end_date=end
        if start ==end: 
            month1=str(int(str(end)[4:6]))+'月'
        else :
            month1=str(int(str(start)[4:6]))+'月'+'-'+str(int(str(end)[4:6]))+'月'
        mask= (df6['开票月份'] >= start_date) & (df6['开票月份'] <= end_date)
        
        
        usetemp=df6.loc[mask]
        e1=usetemp['价税合计'].sum()/100000000
        usetemp1=df6[(df6['开票月份']>=202201)&(df6['开票月份']<=end_date)]
        e2=usetemp1['价税合计'].sum()/100000000
        a1=df.iloc[3,month+8:endmonth+9].sum()  #交易中心所选月份的实际新增
        a2=df1.iloc[2,month+4:endmonth+8].sum() #商务服务部所选月份的实际新增
        a3=df2.iloc[2,month+4:endmonth+8].sum()  #化工事业部所选月份的实际新增
        a4=df3.iloc[2,month+4:endmonth+8].sum()  #市场开发部所选月份的实际新增
        
        temp=df4[df4['招商机构'].isna()]
        temp=temp[temp['会员性质']=='新增']
        temp=temp[temp['收入类别']=='会员费'] #所有月份的个人新增表
        namelist=list(temp['经办人'])
        partlist=list(temp['经办部门']+temp['经办人'])
        persondict=dict(zip(namelist,partlist))
        temp['辅助月份']=temp['日期'].apply(lambda x :int(x[-3:-1]))
        
        b1=b1=len(temp[(temp['辅助月份']>=month)&(temp['辅助月份']<=endmonth)]) # 当前月份的个人新增表
        temp=temp[(temp['辅助月份']>=month)&(temp['辅助月份']<=endmonth)]
        b2=len(temp[temp['经办部门']=='商务服务部']) 
        b3=len(temp[temp['经办部门']=='市场开发部']) 
        b4=len(temp[temp['经办部门']=='化工事业部']) 
        
        
        
        personsee=temp.groupby(['经办人'])
        percount=personsee.count()['收入类别'] 
      #  max_3=list(set(percount))
    #    max_3.sort(reverse=True)
    #    max_3=max_3[:3]
     #   firstlist=list(percount[percount==max_3[0]].index)
    #    secondlist=list(percount[percount==max_3[1]].index)
  #      thirdlist=list(percount[percount==max_3[2]].index)
        
    #    c1=len(firstlist) #+len(secondlist)   #len(thirdlist)  #个人前3的人数
   #     c2=max_3[0]  #排名第一开发的家数
    #    c3=max_3[1]  #排名第二开发的家数
   #     c4=max_3[2]  #排名第三开发的家数

        d1=df.iloc[14,month+8:endmonth+9].sum() #交易中心所选月份收入合计
        d1=round(d1,2)
        d2=df.iloc[8,month+8:endmonth+9].sum() #交易中心所选月份会员费合计
        d2=round(d2,2)
        d3=df.iloc[11,month+8:endmonth+9].sum() #交易中心所选月份财务咨询费合计
        d3=round(d3,2)
        d4=df.iloc[12,month+8:endmonth+9].sum() #交易中心所选月份税收返税费合计
        d4=round(d4,2)
        
        f1=df1.iloc[13,month+4:endmonth+5].sum() #商务服务部所选月份收入合计
        f1=round(f1,2)
        f2=df1.iloc[7,month+4:endmonth+5].sum() #。。。所选月份会员费合计
        f2=round(f2,2)
        f3=df1.iloc[10,month+4:endmonth+5].sum() #。。。所选月份财务咨询费合计
        f3=round(f3,2)
        f4=df1.iloc[11,month+4:endmonth+5].sum() #。。。所选月份税收返税费合计
        f4=round(f4,2)
        f5=df1.iloc[13,4] #。。。累计收入
        f5=round(f5,2)
        
        g1=df1.iloc[18,month+4:endmonth+5].sum() #商务服务部所选月份成本合计
        g1=round(g1,2)
        g2=df1.iloc[14,month+4:endmonth+5].sum() #。。。所选月份机构返还
        g2=round(g2,2)
        g3=df1.iloc[15,month+4:endmonth+5].sum() #。。。所选月份财务咨询费合计
        g3=round(g3,2)
        g4=df1.iloc[16,month+4:endmonth+5].sum() #。。。所选月份税收返税费合计
        g4=round(g4,2)
        g5=df1.iloc[17,month+4:endmonth+5].sum() #。。。所选月份税收返税费合计
        g5=round(g5,2)
        h1=df5.iloc[19,3] #商务服务部全年毛利预算指标
        h1=round(h1,2)
        h0=df5.iloc[13,3] #商务服务部全年收入预算指标
        h0=round(h0,2)
        h2=df5.iloc[19,month+4:endmonth+5].sum() #商务服务部所选月份毛利预算指标
        h2=round(h2,2)
        h3=df1.iloc[19,month+4:endmonth+5].sum() #商务服务部所选月份实际毛利
        h3=round(h3,2)
        h4=df5.iloc[13,month+4:endmonth+5].sum() #商务服务部所选月份收入预算指标
        h4=round(h4,2)
        h5=df1.iloc[19,4] #商务服务部累计实际毛利
        h5=round(h5,2)
        
        
        
        
        q0=df5.iloc[4,3] #商务服务部全年招商预算
        q1=df5.iloc[2,3] #商务服务部全年招商预算（新增）
        q2=df5.iloc[3,3] #商务服务部全年招商预算（续费）
        
        q3=df5.iloc[4,month+4:endmonth+5].sum()  #商务服务部全年招商预算（新增）
        q4=df5.iloc[2,month+4:endmonth+5].sum()  #商务服务部全年招商预算（续费）
        q5=df5.iloc[3,month+4:endmonth+5].sum()  #商务服务部全年招商预算（续费）
        
        q6=df1.iloc[4,3] #商务服务部全年招商预算
        q7=df1.iloc[2,month+4:endmonth+5].sum()  #商务服务部全年招商预算（新增）
        q8=df1.iloc[3,month+4:endmonth+5].sum()  #商务服务部全年招商预算（续费）
        q9=df1.iloc[4,month+4:endmonth+5].sum()  #商务服务部全年招商预算

        picturedata=pd.DataFrame(data=([d2,d3,d4,d1],[f2,f3,f4,f1]),columns=['会员费','财务咨询费','税收返税费','收入合计'],index=['交易中心','交易服务部']) 
        picturedata=pd.DataFrame(picturedata.values.T, index=picturedata.columns, columns=picturedata.index)
        picturedata2=pd.DataFrame(data=([f1,h4],[h3,h2],[q9,q3]),columns=['实际','指标'],index=['收入','毛利','招商数']) 
        #文档主体
        document.add_heading('一、经营指标完成情况',level = 1)
        p = document.add_paragraph(style='Normal')
        run = p.add_run('1. 招商业务情况')
        p = document.add_paragraph(style='Normal')
        run = p.add_run('（1）线上开户情况')
        p = document.add_paragraph(style='Normal')
        run = p.add_run('交易中心'+year+'年'+month1+'月新增线上开户?户，开办至今线上累计开户?户，其中销户?家 目前累计?户。')
        p = document.add_paragraph(style='Normal')
        run = p.add_run('（2）招商会员落地情况')  
        p = document.add_paragraph(style='Normal')
        run = p.add_run('交易中心'+year+'年'+month1+'月新增注册会员'+str(a1)+'家，其中交易服务部开发会员'+str(a2)+'家'+'占比{:.2%}'.format(a2/a1)+'；化工事业部开发会员'+str(a3)+'家，'+'占比{:.2%}'.format(a3/a1)+'；市场开发部开发会员'+str(a4)+'家，'+'占比{:.2%}'.format(a4/a1)+'。开办至今累计落地企业?家。')
    #    p = document.add_paragraph(style='Normal')
   #     run = p.add_run('（3）个人招商会员落地情况') 
    #    p = document.add_paragraph(style='Normal')
  #      run = p.add_run('交易中心'+year+'年'+month1+'月新增个人招商注册会员'+str(b1)+'家，占总新增注册会员的'+'{:.2%}'.format(b1/a1)+'，其中交易服务部个人招商会员企业'+str(b2)+'家,'+'占比{:.2%}'.format(b2/b1)+'；市场开发部个人招商'+str(b3)+'家，'+'占比{:.2%}'.format(b3/b1)+'；化工事业部个人招商'+str(b4)+'家，'+'占比{:.2%}'.format(b4/b1)+'。') 
  #      p = document.add_paragraph(style='Normal')
  #      run = p.add_run('个人招商前三名共'+str(c1)+'人，')
        a111=0
 #       for i in firstlist:
#           a111+=1
 #           if a111< len(firstlist):          
 #               run = p.add_run(persondict.get(i)+'、')
 #           else:run = p.add_run(persondict.get(i)+'招商'+str(c2)+'家，排名第一；')
        a112=0
#        for i in secondlist:
#            a112+=1
#            if a112< len(secondlist):          
#                run = p.add_run(persondict.get(i)+'、')
 #           else:run = p.add_run(persondict.get(i)+'招商'+str(c3)+'家，排名第二；')
        a113=0
#        for i in thirdlist:
#            a113+=1
#            if a113< len(thirdlist):          
##               run = p.add_run(persondict.get(i)+'、')
#            else:run = p.add_run(persondict.get(i)+'招商'+str(c4)+'家，排名第三。')
        
        p = document.add_paragraph(style='Normal')
        run = p.add_run('（3）开票情况')
        p = document.add_paragraph(style='Normal')
        run = p.add_run('交易中心'+year+'年'+month1+'会员企业开票金额'+'{:.5}'.format(e1)+'亿元，今年累计开票金额'+'{:.5}'.format(e2)+'亿元。')
        p = document.add_paragraph(style='Normal')
        run = p.add_run('（4）收入情况')
        p = document.add_paragraph(style='Normal')
        run = p.add_run('交易中心'+year+'年'+month1+'税后收入'+str(d1)+'万元，其中会员费收入'+str(d2)+' 万元（'+'占比{:.2%}'.format(d2/d1)+'），财务咨询费收入'+str(d3)+'万元（'+'占比{:.2%}'.format(d3/d1)+'），税收返还服务费收入'+str(d4)+'万元（'+'占比{:.2%}'.format(d4/d1)+'）。')
        p = document.add_paragraph(style='Normal')
        run = p.add_run('交易服务部'+'年'+month1+'税后收入'+str(f1)+'万元，占交易中心总收入的'+'{:.2%}'.format(f1/d1)+'。其中会员费收入'+str(f2)+' 万元（'+'占比{:.2%}'.format(f2/f1)+'），财务咨询费收入'+str(f3)+'万元（'+'占比{:.2%}'.format(f3/f1)+'），税收返还服务费收入'+str(f4)+'万元（'+'占比{:.2%}'.format(f4/f1)+'）。')
        p = document.add_paragraph(style='Normal')
        run = p.add_run('（5）成本情况')
        p = document.add_paragraph(style='Normal')
        run = p.add_run('交易服务部'+year+'年'+month1+'成本'+str(g1)+'万元，其中招商机构返还成本'+str(g2)+' 万元（'+'占比{:.2%}'.format(g2/g1)+'），财政服务返成'+str(g3)+'万元（'+'占比{:.2%}'.format(g3/g1)+'），商事服务费'+str(g4)+'万元（'+'占比{:.2%}'.format(g4/g1)+'）;'+'财税代理服务费'+str(g5)+'万元（'+'占比{:.2%}'.format(g5/g1)+'）。')
        p = document.add_paragraph(style='Normal')
        run = p.add_run('（6）毛利完成情况')
        p = document.add_paragraph(style='Normal')
        run = p.add_run('交易服务部'+year+'年'+'全年预算指标'+str(h1)+'万元，'+month1+'月预算指标'+str(h2)+'万元，'+month1+'月实际完成'+str(h3)+'万元，完成预算指标的'+'{:.2%}'.format(h3/h2))
        p = document.add_paragraph(style='Normal')
        run = p.add_run('（二）预算目标的差异及原因分析')        
        p = document.add_paragraph(style='Normal')
        run = p.add_run('1.交易服务部全年收入指标'+str(h0)+'万元，'+month1+'收入指标'+str(h4)+'万元。'+month1+'实际收入'+str(f1)+'万元。'+month1+'指标完成度'+'{:.2%}'.format(f1/h4)+';全年目标指标已完成'+'{:.2%}'.format(f5/h0))
        p = document.add_paragraph(style='Normal')
        run = p.add_run('2.交易服务部全年毛利指标'+str(h1)+'万元，'+month1+'毛利指标'+str(h2)+'万元。'+month1+'实际毛利'+str(h3)+'万元。'+month1+'指标完成度'+'{:.2%}'.format(h3/h2)+';全年目标指标已完成'+'{:.2%}'.format(h5/h1))
        p = document.add_paragraph(style='Normal')
        run = p.add_run('3.交易服务部全年招商指标'+str(q0)+'家，（含新增'+str(q1)+'家和续费'+str(q2)+'家),'+month1+'份招商指标'+'{:.1}'.format(q3)+'家，（含新增'+str(q4)+'家和续费'+str(q5)+'家)'+month1+'实际招商'+str(q9)+'家，（含新增'+str(q7)+'家和续费'+str(q8)+'家),全年目标指标已完成'+'{:.2%}'.format(q6/q0))
        
        filename='d:/analyze program/'+year+month1+'经营指标'+'.docx'
        filename1='d:/analyze program/'+year+month1+'经营指标'+'.xlsx'
        filename2='d:/analyze program/'+year+month1+'经营指标2'+'.xlsx'
        document.save(filename)
        picturedata.to_excel(filename1)
        picturedata2.to_excel(filename2)

   
