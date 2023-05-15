# -*- coding: utf-8 -*-
"""
Created on Mon Sep 26 09:32:36 2022

@author: LV
"""
#海口承诺书


import docx
#from docx import Document
from docx.shared import Pt,RGBColor
#from docx.shared import Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
#import xlrd
#import numpy as np
#import pandas as pd

# df 传入要求是已经处理好的结果 path 为文件生成的位置
def make(name,Date,path):
    
    
    

    document=docx.Document()
    document.styles['Normal'].font.size = Pt(17)
    document.styles['Normal'].font.name = u'仿宋_GB2312'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')  
    
    document.styles['Heading 1'].font.size = Pt(22)
    document.styles['Heading 1'].font.name = u'仿宋_GB2312'
    document.styles['Heading 1'].font.color.rgb = RGBColor(0, 0, 0)
    document.styles['Heading 1']._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋') 
    p = document.add_paragraph(style='Heading 1')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('承诺书')

 #   p2 = document.add_paragraph(style='Normal')
 #   run = p2.add_run(' ')
    p2 = document.add_paragraph(style='Normal')
    run = p2.add_run('    我司与海南国际能源交易中心运营总部有限公司（以下简称“运营总部”）签订了《注册会员服务协议》成为运营总部会员企业，根据协议内容运营总部代我司向江东新区管理局提交专项财政扶持的申请资料，我司承诺：')
    p2 = document.add_paragraph(style='Normal')
    run = p2.add_run('    一、已知晓申报项目各项内容的有关规定；')
    p2 = document.add_paragraph(style='Normal')
    run = p2.add_run('    二、所提交申报项目信息及所提供材料真实、准确，符合法律有关规定；')
    p2 = document.add_paragraph(style='Normal')
    run = p2.add_run('    三、本企业诚信经营、合法纳税，未列入失信名单；')
    p2 = document.add_paragraph(style='Normal')
    run = p2.add_run('    四、申报项目未享受过其他财政资金支持；')
    p2 = document.add_paragraph(style='Normal')
    run = p2.add_run('    五、2022年1月1日至2022年10月31日，我司没有因未满足实质性运营要求被工商、税务部门列入过经营异常；')
    p2 = document.add_paragraph(style='Normal')
    run = p2.add_run('    六、如若出现申报金额与实际纳税情况不一致的，我司将协助江东新区管理局做好奖励资金多退少补工作；')
    p2 = document.add_paragraph(style='Normal')
    run = p2.add_run('    七、本企业如隐瞒真实情况或提供虚假材料将获得项目资金应当予以全额退回，并依法向社会公开且列入企业信用信息共享平台失信名单，承担由此引起的所有法律责任。')
    p2 = document.add_paragraph(style='Normal')
    p2 = document.add_paragraph(style='Normal')
#    p2 = document.add_paragraph(style='Normal')

    p2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
   
    seg='（盖章）'

    run = p2.add_run(name+seg)
    p2 = document.add_paragraph(style='Normal')
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    seg2='日期: _______________'
    run = p2.add_run(seg2+Date)
        
    document.save(path)

